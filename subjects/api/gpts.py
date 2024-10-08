from io import BytesIO
import docx
import logging
import openai
import os
import chardet
import pdfplumber
from dotenv import load_dotenv

load_dotenv()  # This loads the .env file automatically
openai_api_key = os.environ.get('API_KEY')

# openai_api_key = os.getenv('API_KEY')
openai_api_key  = os.getenv('API_KEY')

logging.basicConfig(filename='app.log', level=logging.INFO)
logger = logging.getLogger(__name__)


# Summary Generations from gpt
def generate_summary_from_gpt(content, prompt):
    openai.api_key = openai_api_key
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": f"{prompt}\n\n{content[:2000]}"}
        ],
        max_tokens=3000  # Adjust the max tokens based on the required summary length
    )
    summary = response.choices[0].message['content'].strip()
    return summary, prompt


# Key-Points Generations from gpt
def generate_keypoints_from_gpt(content, prompt):
    openai.api_key = openai_api_key
    # prompt = f"You are the key-points generator , Provide me the keypoints in bollet-points, related to the and everytime provide me unique keypoints : \n\n{content[:7500]}"
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": f"{prompt}\n\n{content[:2000]}"}
        ],
        max_tokens=1000  # Adjust the max tokens based on the required summary length
    )
    summary = response.choices[0].message['content'].strip()
    return summary, prompt


# Quizes Generations from gpt
def generate_quizes_from_gpt(content):
    openai.api_key = openai_api_key
    prompt = f"You are the teacher, Provide me quizzes with a maximum limit of 10 questions, each containing 4 options as MCQs. Indicate the correct option by labeling it A, B, C, or D and answer as 'Correct Answer'. Ensure the quizzes are different each time and challenging, apply all previous requirements on this content:\n\n{content[:7500]}\nPlease ensure your response is concise with no extra text, and format questions as 'Question:', and options as 'A:', 'B:', etc and aswer as 'Correct Answer: 'Correct Option''. please strictly follow this pattern for quiz and no other pattern will be allowed for quiz. Only allowed pattern is of question is 'Question:' and nothing else is allowed and make sure no new line will be in between 'Question:' and question heading"
    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=[
            {"role": "system", "content": "You are a helpful assistant."},
            {"role": "user", "content": prompt}
        ],
        max_tokens=2000  # Adjust the max tokens based on the required summary length
    )
    quiz = response.choices[0].message['content'].strip()
    return quiz, prompt

import os

def read_file_content(file):
    # Check if 'file' is a string (i.e., a file path) and open it
    if isinstance(file, str):
        if os.path.exists(file):
            with open(file, 'rb') as f:
                return process_file(f)
        else:
            logger.error("File does not exist.")
            return None
    else:
        return process_file(file)

def process_file(file):
    # Now 'file' is an open file object, so you can safely access 'file.name'
    file_type = file.name.split('.')[-1].lower()
    if file_type == 'pdf':
        return read_pdf_content(file)
    elif file_type == 'txt':
        return read_text_file_content(file)
    elif file_type in ['doc', 'docx']:
        return read_doc_file_content(file)
    else:
        logger.error("Unsupported file type.")
        return None

def read_doc_file_content(file):
    try:
        doc = docx.Document(file)
        doc_content = [para.text for para in doc.paragraphs]
        content = "\n".join(doc_content)
        return content.replace('\x00', '')  # Remove null bytes
    except Exception as e:
        logger.error(f"Error reading DOC/DOCX file: {e}")
        return None

def read_pdf_content(file):
    try:
        max_pages = 100
        pdf_content = []
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= max_pages:
                    break
                text = page.extract_text()
                if text:
                    pdf_content.append(text)
        content = "\n".join(pdf_content)
        return content.replace('\x00', '')  # Remove null bytes
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        return None

def read_text_file_content(file):
    try:
        content = file.read().decode('utf-8')
        return content.replace('\x00', '')  # Remove null bytes
    except Exception as e:
        logger.error(f"Error reading text file: {e}")
        return None
#======================================================= Files reading and converting into text ===========================================
'''
def read_file_content(file):
    file_type = file.name.split('.')[-1].lower()
    if file_type == 'pdf':
        return read_pdf_content(file)
    elif file_type == 'txt':
        return read_text_file_content(file)
    elif file_type in ['doc', 'docx']:
        return read_doc_file_content(file)
    else:
        logger.error("Unsupported file type.")
        return None

def read_doc_file_content(file):
    try:
        doc = docx.Document(file)
        doc_content = [para.text for para in doc.paragraphs]
        content = "\n".join(doc_content)
        return content.replace('\x00', '')  # Remove null bytes
    except Exception as e:
        logger.error(f"Error reading DOC/DOCX file: {e}")
        return None

def read_pdf_content(file):
    try:
        max_pages = 100
        pdf_content = []
        with pdfplumber.open(file) as pdf:
            for i, page in enumerate(pdf.pages):
                if i >= max_pages:
                    break
                text = page.extract_text()
                if text:
                    pdf_content.append(text)
        content = "\n".join(pdf_content)
        return content.replace('\x00', '')  # Remove null bytes
    except Exception as e:
        logger.error(f"Error reading PDF file: {e}")
        return None

def read_text_file_content(file):
    encodings = ['utf-8', 'latin-1']
    content = None

    for encoding in encodings:
        try:
            file.seek(0)  # Reset file pointer to the beginning
            content = file.read().decode(encoding)
            logger.info(f"Successfully decoded with encoding: {encoding}")
            break
        except UnicodeDecodeError:
            logger.warning(f"Failed to decode with encoding: {encoding}")
            continue

    if content is None:
        file.seek(0)
        raw_data = file.read()
        detected_encoding = chardet.detect(raw_data)['encoding']
        try:
            content = raw_data.decode(detected_encoding)
            logger.info(f"Successfully decoded with detected encoding: {detected_encoding}")
        except Exception as e:
            logger.error(f"Failed to decode with detected encoding: {detected_encoding}, Error: {e}")

    if content is None:
        logger.error("Unable to decode file content with any encoding")
        return None

    return content.replace('\x00', '')
'''
